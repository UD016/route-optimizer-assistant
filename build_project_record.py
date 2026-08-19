from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('service_assistant_project_record.docx')

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in('w:tcW')
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_run(run, size=11, bold=False, color='000000', italic=False):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text='', style=None, before=0, after=8, line=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run(r)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_run(r, {1:20, 2:16, 3:14}[level], color='000000')
    return p

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15
for level, size, before, after in [(1,20,20,6),(2,16,18,6),(3,14,16,4)]:
    st = styles[f'Heading {level}']
    st.font.name = 'Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.font.size = Pt(size)
    st.font.bold = False
    st.font.color.rgb = RGBColor(0,0,0)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.15

# Simple Google Docs-style opening block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('Service Assistant Project Record')
set_run(r, 26)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run('Route Optimizer service-department assistant | As of August 14, 2026')
set_run(r, 11, color='555555')

add_para(doc, 'Purpose: preserve a practical record of what has been built, why it was built, how the assistant is organized, and what remains known or unfinished. This document is based on the repository contents, project documentation, and Git history available on August 14, 2026.', after=12)

add_heading(doc, '1. Project at a glance', 1)
add_para(doc, 'The Service Assistant is an embedded AI assistant for the Service department inside the Route Optimizer Streamlit application. It combines conversational reasoning with a curated internal knowledge base, semantic retrieval, technician profiles, temporary file analysis, conversation memory, and a structured dispatch-prioritization tool.', after=8)

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
rows = [
    ('Project repository', 'route-optimizer-assistant'),
    ('Service assistant entry point', 'service_assistant.py'),
    ('Main application', 'app.py'),
    ('Prioritization module', 'prioritization_engine.py'),
    ('Active knowledge base', 'knowledge_base/ (89 files at the time of review)'),
    ('Pilot status', 'Pilot v1.0.0 launched August 13, 2026; end date TBD'),
    ('Current Git revision', '65f6a80 — Update knowledge base and service assistant'),
]
for label, value in rows:
    cells = table.add_row().cells
    cells[0].text = ''
    cells[1].text = ''
    r = cells[0].paragraphs[0].add_run(label); set_run(r, 10, bold=True)
    r = cells[1].paragraphs[0].add_run(value); set_run(r, 10)
set_table_geometry(table, [2700, 6660])

add_heading(doc, '2. Development timeline', 1)
timeline = [
    ('July 13, 2026', 'Initial Service Assistant integration', 'Added the assistant to the Route Optimizer application and introduced the first active knowledge-base structure.'),
    ('July 14–15', 'Knowledge-base organization', 'Added acronyms, renamed and reorganized the knowledge base, and moved from a single-master-document approach toward structured Markdown sources.'),
    ('July 16', 'Conversation memory', 'Implemented persistent multi-turn session memory with a clear-conversation capability using the Agents SDK SQLite session store.'),
    ('July 17–24', 'Operational knowledge expansion', 'Expanded technician profiles and selection logic; added Hydro-Québec invoicing, service rates, overtime entry, cash-customer workflows, generator procedures, CN Rail generator operation, service resources, and software guidance.'),
    ('July 20', 'Semantic retrieval', 'Replaced keyword-oriented retrieval with embedding-based retrieval over Markdown chunks, with a cached embedding index invalidated when knowledge files change.'),
    ('July 21–22', 'Safety and software content', 'Added stop-work authority and restructured the knowledge base into policies, procedures, customers, software, reference, and technician sections.'),
    ('July 27', 'Order-entry refinements', 'Updated cash-customer order-entry and deposit materials and related assistant knowledge.'),
    ('July 29', 'Image and document analysis', 'Added image/PDF analysis and manager setup documentation. Uploaded files are handled as temporary conversation context and are not added to permanent knowledge.'),
    ('August 6–7', 'Prioritization system', 'Added the prioritization knowledge set and a structured engine covering priority levels, customer classes, alarm classification, PM type, comparison, explanations, and scheduling codes.'),
    ('August 10–14', 'Pilot readiness and maintenance', 'Cleaned up source files, added overtime assignment/tracking material, finalized pilot documentation, archived a v1.0.0 assistant snapshot, stopped tracking the local SQLite session database, and refreshed the active knowledge base.'),
]
for date, title, desc in timeline:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(date + ' — '); set_run(r, 11, bold=True)
    r = p.add_run(title + '. '); set_run(r, 11, bold=True)
    r = p.add_run(desc); set_run(r, 11)

add_heading(doc, '3. Current assistant capabilities', 1)
for item in [
    'Conversational answers grounded in the internal Service knowledge base.',
    'Semantic retrieval of the most relevant Markdown excerpts rather than injecting the entire knowledge base into every request.',
    'Technician lookup and recommendation using profiles, capability data, location context, and selection rules.',
    'Structured dispatch prioritization using P1–P6 priority levels, C1–C4 customer classes, alarm classification, PM type, and explicit explanations.',
    'Image analysis and scanned-PDF interpretation through temporary vision inputs; text PDFs and text-based files can also be extracted for the current conversation.',
    'Persistent conversation sessions with the ability to clear or start a new conversation.',
    'Temporary notes and uploaded-file context that remain separate from the permanent knowledge base.',
    'English and French interaction, with the pilot documentation noting that English is generally stronger than French.',
]:
    add_bullet(doc, item)

add_heading(doc, '4. How the service assistant works', 1)
add_heading(doc, 'Knowledge retrieval', 2)
add_para(doc, 'At startup or when needed, service_assistant.py scans Markdown files under knowledge_base/, splits them into chunks, creates embeddings with text-embedding-3-large, and stores the resulting index in .cache/service_assistant_embedding_index.pkl. The cache uses file signatures so edits to the knowledge base trigger a rebuild. For each question, the assistant retrieves relevant chunks and places only those excerpts into the agent context.', after=8)
add_heading(doc, 'Conversation and temporary files', 2)
add_para(doc, 'Session memory is keyed by session ID and backed by service_assistant_sessions.sqlite3 at runtime. The database is intentionally not tracked in Git because it contains local conversation state. Uploaded images, scanned PDFs, and text files are interpreted as temporary context for the current interaction and are not written into knowledge_base/.', after=8)
add_heading(doc, 'Application integration', 2)
add_para(doc, 'app.py hosts the Streamlit experience and integrates the assistant with the broader Route Optimizer application. The application also includes route optimization, technician and depot context, Geotab-related position lookup when configured, the monthly planning page, and timesheet functionality. The Service Assistant is one component within that larger operational tool.', after=8)

add_heading(doc, '5. Knowledge base and source organization', 1)
add_para(doc, 'The active knowledge base contains 89 files organized around the types of information the Service department needs to retrieve quickly:', after=8)
for item in [
    'Customers: Hydro-Québec invoicing and CN Rail generator operation.',
    'Policies: stop-work authority, on-call coverage, working hours, night-work compensation, and absence/vacation management.',
    'Prioritization: governance, decision process, priority levels, customer classification, alarm classification, comparison rules, and FieldAware indications.',
    'Procedures: service calls, quotes, invoices, overtime assignment/tracking, cash-customer workflows, generator fault resets, and related operational processes.',
    'Reference: service rates, department resources, FAQ, tribal knowledge, and acronyms.',
    'Software: BMS, Clover, Power BI, system overview, keyboard shortcuts, favorites, searching, work-order creation, and overtime entry.',
    'Technicians: selection rules, capability matrix, index, and individual technician profiles.',
]:
    add_bullet(doc, item)
add_para(doc, 'An older set of source documents is retained under knowledge_archive/ for reference, including the original service-department master, older technician profiles, hourly-rate files, dispatch priorities, and an older generator fault-reset procedure.', after=8)

add_heading(doc, '6. Pilot record', 1)
add_para(doc, 'Pilot v1.0.0 began August 13, 2026. The initial demonstration received positive reactions: CI690 expressed interest in bringing the platform to the Pointe-Claire branch, and BB44Y asked when it would become available because it could help accelerate his learning. At the time of this record, the pilot end date and knowledge-base version are still open.', after=8)
add_heading(doc, 'Launch limitations recorded', 2)
for item in [
    'No autonomous actions, automatic dispatch, direct BMS actions, live technician availability, or work-order creation.',
    'Image analysis requires a file upload; direct copy/paste of images into the chat was not supported in the manager guide.',
    'Image-heavy conversations can increase token cost, so no-longer-needed images should be removed or the conversation restarted.',
    'Answer quality depends on the quality and coverage of the knowledge base and on the context supplied in the question.',
]:
    add_bullet(doc, item)
add_heading(doc, 'Pilot documentation now in place', 2)
for item in [
    'user_documentation/pilot_log.md: chronological events, decisions, changes, metrics, and end-of-pilot findings.',
    'user_documentation/pilot_feedback_tracking.md: categorized feedback for knowledge gaps, retrieval failures, reasoning issues, data quality, UX, bugs, feature requests, and successful use cases.',
    'user_documentation/pour_patrice.md and its PDF: manager setup, launch, knowledge-base maintenance, and practical pilot guidance.',
]:
    add_bullet(doc, item)

add_heading(doc, '7. Current state and open follow-up items', 1)
add_para(doc, 'The project is at an internal-pilot stage rather than a production-ready autonomous dispatch stage. The core retrieval and operational-assistant foundation is in place, and the immediate next learning loop is pilot feedback.', after=8)
for item in [
    'Complete the pilot log with concrete questions, successful use cases, feedback items, and outcome metrics.',
    'Version or record the active knowledge-base release used by the pilot.',
    'Convert real pilot failures into regression test cases and preserve the original user questions when issues are fixed.',
    'Prioritize knowledge gaps and retrieval failures before adding broad new features.',
    'Decide whether the assistant should remain embedded in Route Optimizer or be made available through another access method.',
    'Evaluate future integrations separately: live technician availability, work-order creation, BMS actions, dispatch automation, and direct image paste support.',
]:
    add_bullet(doc, item)

add_heading(doc, '8. Key files for future maintenance', 1)
for item in [
    'service_assistant.py — retrieval, temporary uploads, vision inputs, session memory, agent construction, and assistant entry point.',
    'app.py — Streamlit application shell and Service Assistant integration.',
    'prioritization_engine.py — structured priority calculations, comparisons, explanations, and helper logic.',
    'knowledge_base/ — active source of truth for internal procedures and reference information.',
    'archives/ — snapshots and legacy assistant/source material retained for historical reference.',
    'user_documentation/ — pilot, feedback, and manager-facing operating documentation.',
    'requirements.txt and runtime.txt — runtime and dependency declarations.',
]:
    add_bullet(doc, item)

add_heading(doc, 'Appendix: repository evidence used', 1)
add_para(doc, 'This record was assembled from the project files and Git history available locally. The service-assistant-specific development sequence begins with commit fd33045 on July 13, 2026 and continues through commit 65f6a80 on August 14, 2026. The repository also contains earlier Route Optimizer work dating back to October 2025.', after=8)
add_para(doc, 'Important scope note: this is a historical record of the implementation visible in the repository, not a claim that every feature has been formally validated in production. Pilot validation and feedback tracking remain ongoing.', after=8)

doc.core_properties.title = 'Service Assistant Project Record'
doc.core_properties.subject = 'Internal project history and pilot record'
doc.core_properties.author = 'Route Optimizer project team'
doc.core_properties.comments = 'Prepared from repository contents and Git history as of 2026-08-14.'
doc.save(OUT)
print(OUT.resolve())
